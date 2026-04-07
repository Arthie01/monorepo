import io
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy import func
from sqlalchemy.orm import Session

from app.data.db import get_db
from app.data.autoparte import Autoparte
from app.data.pedido import Pedido
from app.data.detalle_pedido import DetallePedido
from app.data.usuario_externo import UsuarioExterno
from app.security.auth import verificar_peticion

# ── reportlab ──
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

# ── openpyxl ──
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

# ── python-docx ──
from docx import Document
from docx.shared import RGBColor

# ──────────────────────────────────────────────────────────────────────────────
router = APIRouter(
    prefix="/v1/reportes",
    tags=["Reportes"],
    dependencies=[Depends(verificar_peticion)]   # HTTPBasic en todos los endpoints
)

FORMATOS_VALIDOS = ["pdf", "xlsx", "docx"]
MIME = {
    "pdf":  "application/pdf",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

_TITULOS = {
    "ventas":     "Reporte de Ventas",
    "inventario": "Reporte de Inventario",
    "pedidos":    "Reporte de Pedidos",
    "usuarios":   "Reporte de Usuarios Externos",
}

# colores MACUIN para reportlab
_RED   = colors.HexColor("#C41230")
_DARK  = colors.HexColor("#0D0D0D")
_LIGHT = colors.HexColor("#F5F5F0")

_TS = TableStyle([
    ("BACKGROUND",    (0, 0), (-1,  0), _RED),
    ("TEXTCOLOR",     (0, 0), (-1,  0), colors.white),
    ("FONTNAME",      (0, 0), (-1,  0), "Helvetica-Bold"),
    ("FONTSIZE",      (0, 0), (-1,  0), 9),
    ("FONTSIZE",      (0, 1), (-1, -1), 8),
    ("ROWBACKGROUNDS",(0, 1), (-1, -1), [_LIGHT, colors.white]),
    ("GRID",          (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
    ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
    ("TOPPADDING",    (0, 0), (-1, -1), 4),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
])


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS — CONSULTAS DE DATOS
# ══════════════════════════════════════════════════════════════════════════════

def _date_filter(query, fecha_inicio: Optional[str], fecha_fin: Optional[str]):
    """Aplica filtro de rango de fechas sobre Pedido.creado_en."""
    if fecha_inicio:
        query = query.filter(Pedido.creado_en >= datetime.fromisoformat(fecha_inicio))
    if fecha_fin:
        query = query.filter(Pedido.creado_en <= datetime.fromisoformat(fecha_fin + "T23:59:59"))
    return query


def _datos_ventas(db: Session, fecha_inicio: Optional[str], fecha_fin: Optional[str]) -> dict:
    # KPIs — pedidos no cancelados
    base = _date_filter(
        db.query(Pedido).filter(Pedido.estado != "Cancelado"),
        fecha_inicio, fecha_fin
    )
    pedidos_list = base.all()
    ventas_totales     = sum(float(p.total) for p in pedidos_list)
    pedidos_completados = sum(1 for p in pedidos_list if p.estado == "Completado")

    # Ventas por mes
    mes_q = _date_filter(
        db.query(
            func.to_char(Pedido.creado_en, "YYYY-MM").label("mes"),
            func.sum(Pedido.total).label("total")
        ).filter(Pedido.estado != "Cancelado"),
        fecha_inicio, fecha_fin
    )
    ventas_por_mes = mes_q.group_by("mes").order_by("mes").all()

    # Ventas por categoría
    cat_q = _date_filter(
        db.query(
            Autoparte.categoria,
            func.sum(DetallePedido.precio_unitario * DetallePedido.cantidad).label("total")
        ).join(DetallePedido, Autoparte.id == DetallePedido.autoparte_id)
         .join(Pedido, DetallePedido.pedido_id == Pedido.id)
         .filter(Pedido.estado != "Cancelado"),
        fecha_inicio, fecha_fin
    )
    por_categoria = cat_q.group_by(Autoparte.categoria).order_by(
        func.sum(DetallePedido.precio_unitario * DetallePedido.cantidad).desc()
    ).all()

    # Top 5 productos más vendidos (por revenue)
    top_q = _date_filter(
        db.query(
            Autoparte.nombre, Autoparte.sku, Autoparte.categoria,
            func.sum(DetallePedido.precio_unitario * DetallePedido.cantidad).label("total_vendido"),
            func.sum(DetallePedido.cantidad).label("cantidad_vendida")
        ).join(DetallePedido, Autoparte.id == DetallePedido.autoparte_id)
         .join(Pedido, DetallePedido.pedido_id == Pedido.id)
         .filter(Pedido.estado != "Cancelado"),
        fecha_inicio, fecha_fin
    )
    top_productos = top_q.group_by(
        Autoparte.id, Autoparte.nombre, Autoparte.sku, Autoparte.categoria
    ).order_by(
        func.sum(DetallePedido.precio_unitario * DetallePedido.cantidad).desc()
    ).limit(5).all()

    # Ventas por marca
    marca_q = _date_filter(
        db.query(
            Autoparte.marca,
            func.sum(DetallePedido.precio_unitario * DetallePedido.cantidad).label("total")
        ).join(DetallePedido, Autoparte.id == DetallePedido.autoparte_id)
         .join(Pedido, DetallePedido.pedido_id == Pedido.id)
         .filter(Pedido.estado != "Cancelado"),
        fecha_inicio, fecha_fin
    )
    por_marca = marca_q.group_by(Autoparte.marca).order_by(
        func.sum(DetallePedido.precio_unitario * DetallePedido.cantidad).desc()
    ).all()

    return {
        "periodo": {"inicio": fecha_inicio or "Todos", "fin": fecha_fin or "Todos"},
        "kpis": {
            "ventas_totales":      round(ventas_totales, 2),
            "pedidos_completados": pedidos_completados,
            "total_pedidos":       len(pedidos_list),
        },
        "ventas_por_mes":  [{"mes": r.mes, "total": float(r.total)} for r in ventas_por_mes],
        "por_categoria":   [{"categoria": r.categoria, "total": float(r.total)} for r in por_categoria],
        "top_productos":   [
            {"nombre": r.nombre, "sku": r.sku, "categoria": r.categoria,
             "total_vendido": float(r.total_vendido), "cantidad_vendida": int(r.cantidad_vendida)}
            for r in top_productos
        ],
        "por_marca":       [{"marca": r.marca or "Sin marca", "total": float(r.total)} for r in por_marca],
    }


def _datos_inventario(db: Session) -> dict:
    autopartes = db.query(Autoparte).filter(Autoparte.activo == True).all()

    cat_dict: dict = {}
    for a in autopartes:
        cat_dict[a.categoria] = cat_dict.get(a.categoria, 0) + a.stock
    por_categoria = sorted(
        [{"categoria": k, "stock": v} for k, v in cat_dict.items()],
        key=lambda x: x["stock"], reverse=True
    )

    alertas = sorted(
        [{"nombre": a.nombre, "sku": a.sku, "categoria": a.categoria,
          "stock": a.stock, "stock_minimo": a.stock_minimo}
         for a in autopartes if a.stock < a.stock_minimo],
        key=lambda x: x["stock_minimo"] - x["stock"], reverse=True
    )

    return {
        "generado_en":  datetime.now().strftime("%Y-%m-%d %H:%M"),
        "total_skus":   len(autopartes),
        "stock_total":  sum(a.stock for a in autopartes),
        "por_categoria": por_categoria,
        "alertas":       alertas,
    }


def _datos_pedidos(
    db: Session,
    fecha_inicio: Optional[str],
    fecha_fin: Optional[str],
    estado: Optional[str]
) -> dict:
    q = db.query(Pedido, UsuarioExterno).join(
        UsuarioExterno, Pedido.usuario_externo_id == UsuarioExterno.id, isouter=True
    )
    if fecha_inicio:
        q = q.filter(Pedido.creado_en >= datetime.fromisoformat(fecha_inicio))
    if fecha_fin:
        q = q.filter(Pedido.creado_en <= datetime.fromisoformat(fecha_fin + "T23:59:59"))
    if estado:
        q = q.filter(Pedido.estado.ilike(estado))
    resultados = q.order_by(Pedido.creado_en.desc()).all()

    total_facturado = sum(float(p.total) for p, _ in resultados)

    pedidos_lista = [
        {
            "folio":   p.folio,
            "cliente": f"{u.nombre} {u.apellidos}" if u else "N/A",
            "fecha":   p.creado_en.strftime("%d/%m/%Y") if p.creado_en else "—",
            "estado":  p.estado,
            "total":   float(p.total),
        }
        for p, u in resultados
    ]

    estado_dict: dict = {}
    for p, _ in resultados:
        estado_dict[p.estado] = estado_dict.get(p.estado, 0) + 1
    total_cnt = len(resultados) or 1
    por_estado = sorted(
        [{"estado": k, "cantidad": v, "porcentaje": round(v / total_cnt * 100, 1)}
         for k, v in estado_dict.items()],
        key=lambda x: x["cantidad"], reverse=True
    )

    return {
        "periodo":        {"inicio": fecha_inicio or "Todos", "fin": fecha_fin or "Todos"},
        "filtro_estado":  estado,
        "total_pedidos":  len(resultados),
        "total_facturado": round(total_facturado, 2),
        "pedidos":        pedidos_lista,
        "por_estado":     por_estado,
    }


def _datos_usuarios(db: Session) -> dict:
    resultados = db.query(
        UsuarioExterno,
        func.count(Pedido.id).label("total_pedidos"),
        func.coalesce(func.sum(Pedido.total), 0).label("monto_total"),
    ).outerjoin(Pedido, UsuarioExterno.id == Pedido.usuario_externo_id
    ).group_by(UsuarioExterno.id
    ).order_by(func.coalesce(func.sum(Pedido.total), 0).desc()
    ).all()

    usuarios = [
        {
            "nombre":       f"{u.nombre} {u.apellidos}",
            "empresa":      u.empresa or "—",
            "tipo_cliente": u.tipo_cliente,
            "estado":       u.estado,
            "email":        u.email,
            "creado_en":    u.creado_en.strftime("%d/%m/%Y") if u.creado_en else "—",
            "total_pedidos": int(total_pedidos),
            "monto_total":   float(monto_total),
        }
        for u, total_pedidos, monto_total in resultados
    ]

    return {
        "generado_en":    datetime.now().strftime("%Y-%m-%d %H:%M"),
        "total_usuarios": len(usuarios),
        "usuarios":       usuarios,
    }


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS — GENERADORES DE ARCHIVO
# ══════════════════════════════════════════════════════════════════════════════

def _generar_pdf(datos: dict, tipo: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                             leftMargin=2*cm, rightMargin=2*cm,
                             topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    h1  = ParagraphStyle("h1",  parent=styles["Heading1"],
                          textColor=_RED,  fontSize=18, spaceAfter=4)
    h2  = ParagraphStyle("h2",  parent=styles["Heading2"],
                          textColor=_DARK, fontSize=11, spaceBefore=10, spaceAfter=4)
    sub = ParagraphStyle("sub", parent=styles["Normal"],
                          textColor=colors.HexColor("#6B7280"), fontSize=9, spaceAfter=8)
    ftr = ParagraphStyle("ftr", parent=styles["Normal"],
                          textColor=colors.grey, fontSize=8)

    def tbl(data):
        t = Table(data, hAlign="LEFT")
        t.setStyle(_TS)
        return t

    story = [Paragraph(f"MACUIN — {_TITULOS[tipo]}", h1)]

    if tipo == "ventas":
        story += [Paragraph(f"Periodo: {datos['periodo']['inicio']} — {datos['periodo']['fin']}", sub)]
        k = datos["kpis"]
        story += [Paragraph("KPIs", h2),
                  tbl([["Ventas Totales ($)", "Pedidos Completados", "Total Pedidos"],
                        [f"${k['ventas_totales']:,.2f}", str(k['pedidos_completados']), str(k['total_pedidos'])]])]
        if datos["top_productos"]:
            story += [Paragraph("Top 5 Productos", h2),
                      tbl([["Producto", "SKU", "Categoría", "Vendido ($)", "Cantidad"]] +
                          [[p["nombre"], p["sku"], p["categoria"],
                            f"${p['total_vendido']:,.2f}", str(p["cantidad_vendida"])]
                           for p in datos["top_productos"]])]
        if datos["por_categoria"]:
            story += [Paragraph("Ventas por Categoría", h2),
                      tbl([["Categoría", "Total ($)"]] +
                          [[c["categoria"], f"${c['total']:,.2f}"] for c in datos["por_categoria"]])]
        if datos["por_marca"]:
            story += [Paragraph("Ventas por Marca", h2),
                      tbl([["Marca", "Total ($)"]] +
                          [[m["marca"], f"${m['total']:,.2f}"] for m in datos["por_marca"]])]
        if datos["ventas_por_mes"]:
            story += [Paragraph("Ventas por Mes", h2),
                      tbl([["Mes", "Total ($)"]] +
                          [[m["mes"], f"${m['total']:,.2f}"] for m in datos["ventas_por_mes"]])]

    elif tipo == "inventario":
        story += [Paragraph(
            f"Generado: {datos['generado_en']} · SKUs: {datos['total_skus']} · Stock total: {datos['stock_total']}",
            sub)]
        if datos["por_categoria"]:
            story += [Paragraph("Stock por Categoría", h2),
                      tbl([["Categoría", "Stock"]] +
                          [[c["categoria"], str(c["stock"])] for c in datos["por_categoria"]])]
        if datos["alertas"]:
            story += [Paragraph("Alertas de Stock Bajo", h2),
                      tbl([["Producto", "SKU", "Categoría", "Stock", "Mínimo"]] +
                          [[a["nombre"], a["sku"], a["categoria"],
                            str(a["stock"]), str(a["stock_minimo"])] for a in datos["alertas"]])]

    elif tipo == "pedidos":
        story += [Paragraph(f"Periodo: {datos['periodo']['inicio']} — {datos['periodo']['fin']}", sub)]
        if datos["filtro_estado"]:
            story.append(Paragraph(f"Estado filtrado: {datos['filtro_estado']}", sub))
        story.append(Paragraph(
            f"Total pedidos: {datos['total_pedidos']} · Facturado: ${datos['total_facturado']:,.2f}", sub))
        if datos["por_estado"]:
            story += [Paragraph("Distribución por Estado", h2),
                      tbl([["Estado", "Cantidad", "Porcentaje"]] +
                          [[e["estado"], str(e["cantidad"]), f"{e['porcentaje']}%"]
                           for e in datos["por_estado"]])]
        if datos["pedidos"]:
            story += [Paragraph("Listado de Pedidos", h2),
                      tbl([["Folio", "Cliente", "Fecha", "Estado", "Total ($)"]] +
                          [[p["folio"], p["cliente"], p["fecha"], p["estado"],
                            f"${p['total']:,.2f}"] for p in datos["pedidos"]])]

    elif tipo == "usuarios":
        story += [Paragraph(
            f"Generado: {datos['generado_en']} · Total usuarios: {datos['total_usuarios']}", sub)]
        if datos["usuarios"]:
            story += [Paragraph("Usuarios Externos", h2),
                      tbl([["Nombre", "Empresa", "Tipo Cliente", "Estado", "Pedidos", "Compras ($)"]] +
                          [[u["nombre"], u["empresa"], u["tipo_cliente"], u["estado"],
                            str(u["total_pedidos"]), f"${u['monto_total']:,.2f}"]
                           for u in datos["usuarios"]])]

    story += [Spacer(1, 0.8*cm),
              Paragraph(f"Generado por MACUIN API · {datetime.now().strftime('%d/%m/%Y %H:%M')}", ftr)]
    doc.build(story)
    buffer.seek(0)
    return buffer


def _generar_xlsx(datos: dict, tipo: str) -> io.BytesIO:
    wb = Workbook()
    ws = wb.active
    ws.title = _TITULOS[tipo][:31]

    RED_FILL   = PatternFill("solid", fgColor="C41230")
    LIGHT_FILL = PatternFill("solid", fgColor="F5F5F0")
    WF  = Font(color="FFFFFF", bold=True)
    BF  = Font(bold=True)
    TF  = Font(bold=True, color="C41230", size=14)

    row = [1]   # mutable para closures

    def cell(c, val, font=None, fill=None, align=None):
        ce = ws.cell(row[0], c, val)
        if font:  ce.font = font
        if fill:  ce.fill = fill
        if align: ce.alignment = align
        return ce

    def next_row(n=1):
        row[0] += n

    def section(title):
        next_row()
        ws.cell(row[0], 1, title).font = BF
        next_row()

    def write_table(headers, rows_data):
        for c, h in enumerate(headers, 1):
            ce = ws.cell(row[0], c, h)
            ce.fill = RED_FILL
            ce.font = WF
            ce.alignment = Alignment(horizontal="center")
        next_row()
        for i, r in enumerate(rows_data):
            fill = LIGHT_FILL if i % 2 == 0 else PatternFill()
            for c, v in enumerate(r, 1):
                ce = ws.cell(row[0], c, v)
                ce.fill = fill
            next_row()
        next_row()   # gap

    ws.cell(row[0], 1, f"MACUIN — {_TITULOS[tipo]}").font = TF
    next_row()

    if tipo == "ventas":
        ws.cell(row[0], 1, f"Periodo: {datos['periodo']['inicio']} — {datos['periodo']['fin']}")
        next_row(2)
        section("KPIs")
        k = datos["kpis"]
        write_table(["Ventas Totales ($)", "Pedidos Completados", "Total Pedidos"],
                    [[k["ventas_totales"], k["pedidos_completados"], k["total_pedidos"]]])
        if datos["top_productos"]:
            section("Top 5 Productos")
            write_table(["Producto", "SKU", "Categoría", "Total Vendido ($)", "Cantidad"],
                        [[p["nombre"], p["sku"], p["categoria"],
                          p["total_vendido"], p["cantidad_vendida"]]
                         for p in datos["top_productos"]])
        if datos["por_categoria"]:
            section("Ventas por Categoría")
            write_table(["Categoría", "Total ($)"],
                        [[c["categoria"], c["total"]] for c in datos["por_categoria"]])
        if datos["por_marca"]:
            section("Ventas por Marca")
            write_table(["Marca", "Total ($)"],
                        [[m["marca"], m["total"]] for m in datos["por_marca"]])
        if datos["ventas_por_mes"]:
            section("Ventas por Mes")
            write_table(["Mes", "Total ($)"],
                        [[m["mes"], m["total"]] for m in datos["ventas_por_mes"]])

    elif tipo == "inventario":
        ws.cell(row[0], 1,
                f"Generado: {datos['generado_en']} | SKUs: {datos['total_skus']} | Stock Total: {datos['stock_total']}")
        next_row(2)
        if datos["por_categoria"]:
            section("Stock por Categoría")
            write_table(["Categoría", "Stock"],
                        [[c["categoria"], c["stock"]] for c in datos["por_categoria"]])
        if datos["alertas"]:
            section("Alertas de Stock Bajo")
            write_table(["Producto", "SKU", "Categoría", "Stock Actual", "Stock Mínimo"],
                        [[a["nombre"], a["sku"], a["categoria"],
                          a["stock"], a["stock_minimo"]] for a in datos["alertas"]])

    elif tipo == "pedidos":
        ws.cell(row[0], 1, f"Periodo: {datos['periodo']['inicio']} — {datos['periodo']['fin']}")
        next_row()
        if datos["filtro_estado"]:
            ws.cell(row[0], 1, f"Estado filtrado: {datos['filtro_estado']}")
            next_row()
        ws.cell(row[0], 1, f"Total: {datos['total_pedidos']} pedidos | ${datos['total_facturado']:,.2f}")
        next_row(2)
        if datos["por_estado"]:
            section("Distribución por Estado")
            write_table(["Estado", "Cantidad", "Porcentaje (%)"],
                        [[e["estado"], e["cantidad"], e["porcentaje"]] for e in datos["por_estado"]])
        if datos["pedidos"]:
            section("Listado de Pedidos")
            write_table(["Folio", "Cliente", "Fecha", "Estado", "Total ($)"],
                        [[p["folio"], p["cliente"], p["fecha"], p["estado"], p["total"]]
                         for p in datos["pedidos"]])

    elif tipo == "usuarios":
        ws.cell(row[0], 1,
                f"Generado: {datos['generado_en']} | Total: {datos['total_usuarios']} usuarios")
        next_row(2)
        section("Usuarios Externos")
        write_table(
            ["Nombre", "Empresa", "Tipo Cliente", "Estado", "Total Pedidos", "Monto Total ($)"],
            [[u["nombre"], u["empresa"], u["tipo_cliente"], u["estado"],
              u["total_pedidos"], u["monto_total"]] for u in datos["usuarios"]]
        )

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_docx(datos: dict, tipo: str) -> io.BytesIO:
    doc = Document()

    title = doc.add_heading(f"MACUIN — {_TITULOS[tipo]}", 0)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0xC4, 0x12, 0x30)

    def add_tbl(headers, rows_data):
        t = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
        t.style = "Table Grid"
        for c, h in enumerate(headers):
            cell = t.rows[0].cells[c]
            cell.text = h
            if cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0xC4, 0x12, 0x30)
        for r_i, row_data in enumerate(rows_data, 1):
            for c_i, val in enumerate(row_data):
                t.rows[r_i].cells[c_i].text = str(val)
        doc.add_paragraph()

    if tipo == "ventas":
        doc.add_paragraph(f"Periodo: {datos['periodo']['inicio']} — {datos['periodo']['fin']}")
        k = datos["kpis"]
        doc.add_heading("KPIs", 2)
        add_tbl(["Ventas Totales ($)", "Pedidos Completados", "Total Pedidos"],
                [[f"${k['ventas_totales']:,.2f}", k["pedidos_completados"], k["total_pedidos"]]])
        if datos["top_productos"]:
            doc.add_heading("Top 5 Productos", 2)
            add_tbl(["Producto", "SKU", "Categoría", "Total ($)", "Cantidad"],
                    [[p["nombre"], p["sku"], p["categoria"],
                      f"${p['total_vendido']:,.2f}", p["cantidad_vendida"]]
                     for p in datos["top_productos"]])
        if datos["por_categoria"]:
            doc.add_heading("Ventas por Categoría", 2)
            add_tbl(["Categoría", "Total ($)"],
                    [[c["categoria"], f"${c['total']:,.2f}"] for c in datos["por_categoria"]])
        if datos["por_marca"]:
            doc.add_heading("Ventas por Marca", 2)
            add_tbl(["Marca", "Total ($)"],
                    [[m["marca"], f"${m['total']:,.2f}"] for m in datos["por_marca"]])

    elif tipo == "inventario":
        doc.add_paragraph(
            f"Generado: {datos['generado_en']} | SKUs: {datos['total_skus']} | Stock Total: {datos['stock_total']}")
        if datos["por_categoria"]:
            doc.add_heading("Stock por Categoría", 2)
            add_tbl(["Categoría", "Stock"],
                    [[c["categoria"], c["stock"]] for c in datos["por_categoria"]])
        if datos["alertas"]:
            doc.add_heading("Alertas de Stock Bajo", 2)
            add_tbl(["Producto", "SKU", "Categoría", "Stock", "Mínimo"],
                    [[a["nombre"], a["sku"], a["categoria"],
                      a["stock"], a["stock_minimo"]] for a in datos["alertas"]])

    elif tipo == "pedidos":
        doc.add_paragraph(f"Periodo: {datos['periodo']['inicio']} — {datos['periodo']['fin']}")
        if datos["filtro_estado"]:
            doc.add_paragraph(f"Estado filtrado: {datos['filtro_estado']}")
        doc.add_paragraph(f"Total: {datos['total_pedidos']} pedidos | ${datos['total_facturado']:,.2f}")
        if datos["por_estado"]:
            doc.add_heading("Distribución por Estado", 2)
            add_tbl(["Estado", "Cantidad", "Porcentaje (%)"],
                    [[e["estado"], e["cantidad"], f"{e['porcentaje']}%"] for e in datos["por_estado"]])
        if datos["pedidos"]:
            doc.add_heading("Listado de Pedidos", 2)
            add_tbl(["Folio", "Cliente", "Fecha", "Estado", "Total ($)"],
                    [[p["folio"], p["cliente"], p["fecha"], p["estado"],
                      f"${p['total']:,.2f}"] for p in datos["pedidos"]])

    elif tipo == "usuarios":
        doc.add_paragraph(
            f"Generado: {datos['generado_en']} | Total: {datos['total_usuarios']} usuarios")
        if datos["usuarios"]:
            doc.add_heading("Usuarios Externos", 2)
            add_tbl(["Nombre", "Empresa", "Tipo Cliente", "Estado", "Pedidos", "Monto Total ($)"],
                    [[u["nombre"], u["empresa"], u["tipo_cliente"], u["estado"],
                      u["total_pedidos"], f"${u['monto_total']:,.2f}"]
                     for u in datos["usuarios"]])

    doc.add_paragraph(f"Generado por MACUIN API — {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ══════════════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

def _responder(buffer: io.BytesIO, formato: str, nombre_base: str) -> StreamingResponse:
    filename = f"{nombre_base}_{datetime.now().strftime('%Y%m%d')}.{formato}"
    return StreamingResponse(
        buffer,
        media_type=MIME[formato],
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

def _validar_formato(formato: str):
    if formato not in FORMATOS_VALIDOS:
        raise HTTPException(
            status_code=400,
            detail=f"Formato '{formato}' no válido. Opciones: {', '.join(FORMATOS_VALIDOS)}"
        )

def _despachar(datos: dict, tipo: str, formato: str) -> io.BytesIO:
    if formato == "pdf":
        return _generar_pdf(datos, tipo)
    elif formato == "xlsx":
        return _generar_xlsx(datos, tipo)
    else:
        return _generar_docx(datos, tipo)


@router.get("/ventas/{formato}", status_code=status.HTTP_200_OK)
async def reporte_ventas(
    formato: str,
    fecha_inicio: Optional[str] = None,
    fecha_fin:    Optional[str] = None,
    db: Session = Depends(get_db)
):
    _validar_formato(formato)
    datos = _datos_ventas(db, fecha_inicio, fecha_fin)
    return _responder(_despachar(datos, "ventas", formato), formato, "reporte_ventas")


@router.get("/inventario/{formato}", status_code=status.HTTP_200_OK)
async def reporte_inventario(
    formato: str,
    db: Session = Depends(get_db)
):
    _validar_formato(formato)
    datos = _datos_inventario(db)
    return _responder(_despachar(datos, "inventario", formato), formato, "reporte_inventario")


@router.get("/pedidos/{formato}", status_code=status.HTTP_200_OK)
async def reporte_pedidos(
    formato: str,
    fecha_inicio: Optional[str] = None,
    fecha_fin:    Optional[str] = None,
    estado:       Optional[str] = None,
    db: Session = Depends(get_db)
):
    _validar_formato(formato)
    datos = _datos_pedidos(db, fecha_inicio, fecha_fin, estado)
    return _responder(_despachar(datos, "pedidos", formato), formato, "reporte_pedidos")


@router.get("/usuarios/{formato}", status_code=status.HTTP_200_OK)
async def reporte_usuarios(
    formato: str,
    db: Session = Depends(get_db)
):
    _validar_formato(formato)
    datos = _datos_usuarios(db)
    return _responder(_despachar(datos, "usuarios", formato), formato, "reporte_usuarios")
